import datetime
from datetime import timedelta
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
import io
import uuid

# --- CONSTANTS ---
RATES = {
    "Level 2: Coordination of Supports": 100.14,
    "Level 3: Specialist Support Coordination": 190.41
}

# --- CSV HANDLERS ---
def generate_csv_template():
    """Creates a blank CSV template for bulk imports."""
    df = pd.DataFrame(columns=["Name", "NDIS Number", "Support Level", "Total Budget", "Current Balance", "Plan End Date", "Hours Per Week"])
    # Add one example row
    df.loc[0] = ["John Doe", "430123456", "Level 2: Coordination of Supports", 18000, 15000, (datetime.date.today() + timedelta(weeks=40)).strftime("%Y-%m-%d"), 1.5]
    return df.to_csv(index=False).encode('utf-8')

def process_csv_upload(uploaded_file):
    """Converts uploaded CSV into the app's client dictionary format."""
    try:
        df = pd.read_csv(uploaded_file)
        new_clients = []
        
        for _, row in df.iterrows():
            # Flexible Column Matching
            name = row.get("Name") or row.get("Participant") or "Unknown"
            ndis = row.get("NDIS Number") or row.get("NDIS #") or ""
            
            # Level Matching
            raw_level = row.get("Support Level") or "Level 2"
            level = "Level 2: Coordination of Supports" # Default
            if "3" in str(raw_level): 
                level = "Level 3: Specialist Support Coordination"
                
            rate = RATES.get(level, 100.14)
            
            # Clean Numbers (Remove $ or commas)
            def clean_num(val):
                if pd.isna(val): return 0.0
                return float(str(val).replace('$','').replace(',','').strip())

            budget = clean_num(row.get("Total Budget"))
            balance = clean_num(row.get("Current Balance"))
            hours = clean_num(row.get("Hours Per Week") or row.get("Hours"))

            # Date Handling
            date_val = row.get("Plan End Date") or row.get("End Date")
            if not date_val or pd.isna(date_val):
                plan_end = datetime.date.today() + timedelta(weeks=40)
            else:
                plan_end = str(date_val).split(' ')[0] # Keep simple string YYYY-MM-DD

            client = {
                "id": str(uuid.uuid4()),
                "name": str(name),
                "ndis_number": str(ndis),
                "level": level,
                "rate": rate,
                "budget": budget,
                "balance": balance,
                "plan_end": plan_end,
                "hours": hours,
                "notes": ""
            }
            new_clients.append(client)
            
        return new_clients
    except Exception as e:
        st.error(f"CSV Error: {e}")
        return None

# --- MATH ENGINE ---
def calculate_client_metrics(c):
    """Calculates runway, surplus, and status."""
    try:
        balance = float(c.get('balance', 0))
        hours = float(c.get('hours', 0))
        rate = float(c.get('rate', 100.14))
        budget = float(c.get('budget', 0))
        
        # Handle Date Parsing
        plan_end_str = c.get('plan_end')
        if isinstance(plan_end_str, (datetime.date, datetime.datetime)):
            plan_end = plan_end_str
        else:
            try:
                plan_end = datetime.datetime.strptime(str(plan_end_str), "%Y-%m-%d").date()
            except:
                try:
                    # Try DD/MM/YYYY
                    plan_end = datetime.datetime.strptime(str(plan_end_str), "%d/%m/%Y").date()
                except:
                    plan_end = datetime.date.today() + timedelta(weeks=40)
            
    except Exception:
        return None

    today = datetime.date.today()
    weeks_remaining = max(0, (plan_end - today).days / 7)
    weekly_cost = hours * rate
    
    if weekly_cost > 0:
        runway_weeks = balance / weekly_cost
    else:
        runway_weeks = 999 
        
    surplus = balance - (weekly_cost * weeks_remaining)
    depletion_date = today + timedelta(days=int(runway_weeks * 7))
    
    # NDIS Status Logic
    if runway_weeks >= weeks_remaining * 1.2:
        status = "ROBUST SURPLUS"
        color = "#3fb950" # Green
    elif runway_weeks >= weeks_remaining:
        status = "SUSTAINABLE"
        color = "#2ea043" # Light Green
    elif runway_weeks >= max(0, weeks_remaining - 4):
        status = "MONITORING REQUIRED"
        color = "#d29922" # Yellow
    else:
        status = "CRITICAL SHORTFALL"
        color = "#f85149" # Red

    return {
        "id": c.get('id'),
        "name": c.get('name', 'Unknown'),
        "ndis_number": c.get('ndis_number', ''),
        "level": c.get('level'),
        "rate": rate,
        "budget": budget,
        "balance": balance,
        "hours": hours,
        "plan_end": plan_end,
        "weeks_remaining": weeks_remaining,
        "weekly_cost": weekly_cost,
        "runway_weeks": runway_weeks,
        "depletion_date": depletion_date,
        "surplus": surplus,
        "status": status,
        "color": color,
        "notes": c.get('notes', '')
    }

# --- WORD REPORT GENERATOR ---
def generate_caseload_report(caseload_data):
    """Generates a professional Word doc."""
    doc = Document()
    doc.add_heading('XYSTON | Caseload Master Report', 0)
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%d %B %Y')}")
    doc.add_paragraph("Confidential: Internal Use Only")
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    total_funds = sum(c['balance'] for c in caseload_data)
    monthly_rev = sum(c['weekly_cost'] for c in caseload_data) * 4.33
    critical_clients = [c for c in caseload_data if "CRITICAL" in c['status']]
    
    p = doc.add_paragraph()
    p.add_run(f"Total Participants: {len(caseload_data)}\n").bold = True
    p.add_run(f"Funds Under Management: ${total_funds:,.2f}\n")
    p.add_run(f"Projected Monthly Revenue: ${monthly_rev:,.2f}\n")
    
    if critical_clients:
        doc.add_heading('⚠️ Critical Risk Watchlist', 3)
        for c in critical_clients:
            doc.add_paragraph(f"• {c['name']}: Runs out on {c['depletion_date'].strftime('%d/%m/%y')} (${abs(c['surplus']):,.0f} shortfall)")
            
    doc.add_page_break()
    
    # Individual Files
    for c in caseload_data:
        doc.add_heading(f"{c['name']} ({c['ndis_number']})", 1)
        
        p = doc.add_paragraph()
        run = p.add_run(f"PLAN HEALTH: {c['status']}")
        run.bold = True
        run.font.color.rgb = RGBColor(0xf8, 0x51, 0x49) if "CRITICAL" in c['status'] else RGBColor(0x3f, 0xb9, 0x50)
        
        p = doc.add_paragraph()
        p.add_run(f"Current Balance: ${c['balance']:,.2f}\n")
        p.add_run(f"Weekly Burn: ${c['weekly_cost']:,.2f} ({c['hours']} hrs/wk)\n")
        p.add_run(f"Plan Ends: {c['plan_end'].strftime('%d/%m/%Y')} ({c['weeks_remaining']:.1f} wks left)\n")
        
        outcome = f"${c['surplus']:,.2f}"
        p.add_run(f"Projected Outcome: {'+' if c['surplus'] > 0 else ''}{outcome}\n")
        
        if c['notes']:
            doc.add_heading('Strategy Notes', 2)
            doc.add_paragraph(c['notes'])
            
        doc.add_page_break()
        
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
